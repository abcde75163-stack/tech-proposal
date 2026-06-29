import io
import streamlit as st


def extract_text_from_docx(file_bytes: bytes) -> str:
    """DOCX 파일에서 텍스트 추출"""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 표 내용도 추출
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        st.error(f"DOCX 파싱 오류: {e}")
        return ""


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """PDF 파일에서 텍스트 추출"""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        texts = []
        for page in doc:
            texts.append(page.get_text())
        doc.close()
        return "\n".join(texts)
    except Exception as e:
        st.error(f"PDF 파싱 오류: {e}")
        return ""


def parse_uploaded_file(uploaded_file) -> str:
    """업로드된 파일 유형에 따라 자동 파싱"""
    if uploaded_file is None:
        return ""
    
    file_bytes = uploaded_file.read()
    filename = uploaded_file.name.lower()
    
    if filename.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif filename.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    else:
        st.warning("지원하지 않는 파일 형식입니다. DOCX 또는 PDF만 가능합니다.")
        return ""
