import io
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_color: str):
    """표 셀 배경색 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int):
    """제목 추가"""
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p


def parse_markdown_table(lines: list) -> list:
    """Markdown 표를 행/열 리스트로 파싱"""
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith("|") and not re.match(r"^\|[-| ]+\|$", line):
            cells = [c.strip() for c in line.strip("|").split("|")]
            rows.append(cells)
    return rows


def add_table_to_doc(doc: Document, rows: list):
    """파싱된 행을 DOCX 표로 추가"""
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                cell.text = cell_text
                # 첫 행은 헤더 스타일
                if i == 0:
                    set_cell_shading(cell, "D5E8F0")
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                            run.font.size = Pt(9)
                else:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(9)
    doc.add_paragraph()


def markdown_to_docx(markdown_text: str) -> bytes:
    """Markdown 텍스트를 DOCX bytes로 변환"""
    doc = Document()

    # 기본 스타일 설정
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)

    # A4 페이지 설정
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # 제목
        if line.startswith("#### "):
            add_heading(doc, line[5:], 4)
        elif line.startswith("### "):
            add_heading(doc, line[4:], 3)
        elif line.startswith("## "):
            add_heading(doc, line[3:], 2)
        elif line.startswith("# "):
            add_heading(doc, line[2:], 1)

        # 표 감지
        elif line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_markdown_table(table_lines)
            add_table_to_doc(doc, rows)
            continue

        # 인용 블록
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:])
            p.style = "Quote" if "Quote" in [s.name for s in doc.styles] else "Normal"
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                run.italic = True

        # 구분선
        elif line.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)

        # 글머리 기호
        elif line.startswith("- ") or line.startswith("* "):
            text = line[2:]
            p = doc.add_paragraph(style="List Bullet")
            # **bold** 처리
            parts = re.split(r"(\*\*[^*]+\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        # 번호 목록
        elif re.match(r"^\d+\. ", line):
            text = re.sub(r"^\d+\. ", "", line)
            p = doc.add_paragraph(style="List Number")
            p.add_run(text)

        # 빈 줄
        elif line.strip() == "":
            pass

        # 일반 텍스트 (**bold** 처리 포함)
        else:
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*[^*]+\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
